"""
report_generator.py — Turns already-calculated analytics into charts, an
optional Grok narrative, and PDF/DOCX report files.

Pipeline (per the Step 4 spec):
    PostgreSQL -> Python calculations (analytics.py) -> Tables/Charts (here)
    -> Grok narrative (here, constrained) -> Final Report (here)

Grok is used ONLY to phrase already-computed numbers into prose. It is
never asked to calculate anything, and every prompt sent to it embeds the
exact numbers it's allowed to reference plus an explicit instruction not
to invent facts, trends, mines, owners, or sources. If Grok is unavailable
(no API key, no network, or a bad response), a deterministic template
fallback is used instead — built directly from the same numbers with
plain Python string formatting, so the report is never blocked on an LLM.
"""

from __future__ import annotations

import json
import os
from datetime import datetime, timezone
from typing import Any

import matplotlib
matplotlib.use("Agg")  # headless — no display server in this environment
import matplotlib.pyplot as plt

GROK_API_URL = "https://api.x.ai/v1/chat/completions"
GROK_MODEL = "grok-4"

REPORT_SECTIONS = [
    "executive_summary",
    "production_analysis",
    "mine_owner_analysis",
    "key_findings",
    "data_quality",
    "important_observations",
]


# ---------------------------------------------------------------------------
# Charts — matplotlib, rendered straight from the analytics dict
# ---------------------------------------------------------------------------
def _style_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(labelsize=8)


def generate_charts(analytics: dict[str, Any], output_dir: str) -> dict[str, str]:
    """Renders every chart the data supports and returns {name: file_path}."""
    os.makedirs(output_dir, exist_ok=True)
    charts: dict[str, str] = {}

    # Production by state (top 10) — only if any state has production data
    by_state = [r for r in analytics["by_state"] if r["total_production_mt"] > 0][:10]
    if by_state:
        fig, ax = plt.subplots(figsize=(7, 4))
        names = [r["state_ut"] for r in reversed(by_state)]
        values = [r["total_production_mt"] for r in reversed(by_state)]
        ax.barh(names, values, color="#3aa675")
        ax.set_xlabel("Production (MT)")
        ax.set_title("Production by State (Top 10)")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "by_state.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["by_state"] = path

    # Top producing mines
    top_mines = analytics["top_mines"][:10]
    if top_mines:
        fig, ax = plt.subplots(figsize=(7, 4))
        names = [m["mine_name"][:30] for m in reversed(top_mines)]
        values = [m["production_mt"] for m in reversed(top_mines)]
        ax.barh(names, values, color="#2f7fb8")
        ax.set_xlabel("Production (MT)")
        ax.set_title("Top Producing Mines")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "top_mines.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["top_mines"] = path

    # Production by owner (top 10)
    by_owner = [r for r in analytics["by_owner"] if r["total_production_mt"] > 0][:10]
    if by_owner:
        fig, ax = plt.subplots(figsize=(7, 4))
        names = [r["owner"][:30] for r in reversed(by_owner)]
        values = [r["total_production_mt"] for r in reversed(by_owner)]
        ax.barh(names, values, color="#d9a53f")
        ax.set_xlabel("Production (MT)")
        ax.set_title("Production by Owner (Top 10)")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "by_owner.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["by_owner"] = path

    # Coal vs lignite
    cvl = analytics["coal_vs_lignite"]
    if cvl and len(cvl) > 1:
        fig, ax = plt.subplots(figsize=(4.5, 4.5))
        labels = [r["category"] for r in cvl]
        counts = [r["mine_count"] for r in cvl]
        ax.pie(counts, labels=labels, autopct="%1.0f%%", colors=["#3aa675", "#d9a53f", "#8b98a5"])
        ax.set_title("Mine Count: Coal vs Lignite")
        fig.tight_layout()
        path = os.path.join(output_dir, "coal_vs_lignite.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["coal_vs_lignite"] = path

    # Mine type distribution
    mtd = analytics["mine_type_distribution"]
    if mtd:
        fig, ax = plt.subplots(figsize=(6, 4))
        labels = [r["mine_type"] for r in mtd]
        counts = [r["mine_count"] for r in mtd]
        ax.bar(labels, counts, color="#8b6fd9")
        ax.set_ylabel("Mine count")
        ax.set_title("Mine Type Distribution")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "mine_type.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["mine_type"] = path

    # Government vs private
    own = analytics["ownership_distribution"]
    if own:
        fig, ax = plt.subplots(figsize=(6, 4))
        labels = [r["ownership_label"] for r in own]
        counts = [r["mine_count"] for r in own]
        ax.bar(labels, counts, color="#d96f6f")
        ax.set_ylabel("Mine count")
        ax.set_title("Government vs Private Distribution")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "ownership.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["ownership"] = path

    # Validation statistics
    val = analytics["validation"]["by_status"]
    if sum(val.values()) > 0:
        fig, ax = plt.subplots(figsize=(6, 4))
        colors = {"VERIFIED": "#3aa675", "NEEDS_REVIEW": "#d9a53f", "WARNING": "#e69628", "ERROR": "#d9534f"}
        labels = list(val.keys())
        counts = list(val.values())
        ax.bar(labels, counts, color=[colors.get(l, "#8b98a5") for l in labels])
        ax.set_ylabel("Mine count")
        ax.set_title("Validation Status")
        _style_axes(ax)
        fig.tight_layout()
        path = os.path.join(output_dir, "validation.png")
        fig.savefig(path, dpi=150)
        plt.close(fig)
        charts["validation"] = path

    return charts


# ---------------------------------------------------------------------------
# Grok narrative — strictly constrained to the supplied numbers
# ---------------------------------------------------------------------------
def _template_fallback_narrative(analytics: dict[str, Any], document_info: dict[str, Any]) -> dict[str, str]:
    """
    Deterministic, template-based narrative built directly from the same
    numbers — used whenever Grok is unavailable. No numbers here are
    invented; every sentence just restates a value already present in
    `analytics`.
    """
    t = analytics["totals"]
    v = analytics["validation"]
    top_state = analytics["by_state"][0] if analytics["by_state"] else None
    top_mine = analytics["top_mines"][0] if analytics["top_mines"] else None
    top_owner = analytics["by_owner"][0] if analytics["by_owner"] else None

    exec_summary = (
        f"This report covers {t['mine_count']} mine records from the "
        f"{analytics['dataset_period']} Indian Coal Mines dataset "
        f"(source file: {document_info.get('filename', 'the uploaded document')}). "
        f"Total recorded production across mines with a reported figure is "
        f"{t['total_production_mt']:,} MT, from {t['mines_with_production']} of "
        f"{t['mine_count']} mines ({t['mines_missing_production']} have no reported "
        f"production value in the source data)."
    )

    prod_analysis = (
        f"Average production per mine (where reported) is {t['avg_production_mt']:,} MT, "
        f"ranging from {t['min_production_mt']:,} MT to {t['max_production_mt']:,} MT. "
    )
    if top_state:
        prod_analysis += (
            f"{top_state['state_ut']} has the highest recorded total production "
            f"({top_state['total_production_mt']:,} MT across {top_state['mine_count']} mines)."
        )

    mine_owner = ""
    if top_mine:
        mine_owner += (
            f"The single highest-producing mine in the dataset is {top_mine['mine_name']} "
            f"({top_mine['state_ut'] or 'state not recorded'}), with {top_mine['production_mt']:,} MT. "
        )
    if top_owner:
        mine_owner += (
            f"By owner, {top_owner['owner']} accounts for the highest total production "
            f"({top_owner['total_production_mt']:,} MT across {top_owner['mine_count']} mines)."
        )

    key_findings = (
        f"Of {v['total_mines']} mine records, {v['by_status']['VERIFIED']} passed all "
        f"automated validation checks (VERIFIED), {v['by_status']['NEEDS_REVIEW']} are "
        f"flagged NEEDS_REVIEW, {v['by_status']['WARNING']} are flagged WARNING, and "
        f"{v['by_status']['ERROR']} are flagged ERROR. Average confidence across all "
        f"records is {v['avg_confidence']}."
    )

    data_quality = (
        "The most common validation issues are: " +
        "; ".join(f"{i['issue']} ({i['count']} records)" for i in v["top_issue_types"][:5])
        if v["top_issue_types"] else
        "No validation issues were recorded."
    )

    observations = (
        "This dataset represents a single period (2019-2020) with no prior-period data "
        "available for comparison, so no year-over-year trend is reported. "
        f"{t['mines_missing_production']} mines have no reported production figure and "
        "are excluded from production totals and averages."
    )

    return {
        "executive_summary": exec_summary,
        "production_analysis": prod_analysis,
        "mine_owner_analysis": mine_owner or "Insufficient owner/production data to summarize.",
        "key_findings": key_findings,
        "data_quality": data_quality,
        "important_observations": observations,
    }


def _build_grok_prompt(analytics: dict[str, Any], document_info: dict[str, Any]) -> str:
    # Only the calculated numbers are included — Grok sees nothing else and
    # is explicitly told it may not add anything not present here.
    payload = {
        "dataset_period": analytics["dataset_period"],
        "document": {
            "filename": document_info.get("filename"),
            "sheet_name": document_info.get("sheet_name"),
            "total_rows": document_info.get("total_rows"),
        },
        "totals": analytics["totals"],
        "top_5_states_by_production": analytics["by_state"][:5],
        "top_5_owners_by_production": analytics["by_owner"][:5],
        "top_5_mines_by_production": analytics["top_mines"][:5],
        "coal_vs_lignite": analytics["coal_vs_lignite"],
        "mine_type_distribution": analytics["mine_type_distribution"],
        "ownership_distribution": analytics["ownership_distribution"],
        "validation": analytics["validation"],
    }
    return (
        "You are writing narrative sections for a factual data report about Indian coal "
        "mines. Below is the COMPLETE set of numbers you are allowed to reference, already "
        "calculated from PostgreSQL/Python. Do not calculate, estimate, round differently, "
        "invent, or modify any number. Do not invent mines, owners, states, trends, sources, "
        "or facts not present in this JSON. Do not claim any year-over-year trend — this "
        "dataset covers a single period only. If something cannot be determined from this "
        "data, say so explicitly instead of guessing.\n\n"
        f"DATA:\n{json.dumps(payload, indent=2, default=str)}\n\n"
        "Return ONLY a JSON object with exactly these string keys, each containing a "
        "2-5 sentence paragraph in plain, professional report language: "
        f"{', '.join(REPORT_SECTIONS)}."
    )


def generate_narrative(analytics: dict[str, Any], document_info: dict[str, Any]) -> tuple[dict[str, str], str]:
    """
    Returns (sections_dict, narrative_source) where narrative_source is
    either "grok" or "template_fallback" (with a reason appended).
    """
    api_key = os.getenv("GROK_API_KEY") or os.getenv("XAI_API_KEY")
    if not api_key:
        return _template_fallback_narrative(analytics, document_info), "template_fallback (GROK_API_KEY not set)"

    try:
        import requests

        prompt = _build_grok_prompt(analytics, document_info)
        resp = requests.post(
            GROK_API_URL,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": GROK_MODEL,
                "messages": [
                    {"role": "system", "content": "You write factual report narrative from supplied data only."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.2,
            },
            timeout=30,
        )
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        # Strip markdown fences if present
        content = content.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        sections = json.loads(content)
        if not all(k in sections for k in REPORT_SECTIONS):
            raise ValueError("Grok response missing required sections")
        return sections, "grok"
    except Exception as exc:  # noqa: BLE001 — any failure falls back safely
        fallback = _template_fallback_narrative(analytics, document_info)
        return fallback, f"template_fallback (Grok error: {exc})"


# ---------------------------------------------------------------------------
# PDF assembly (reportlab)
# ---------------------------------------------------------------------------
def _pdf_table(data: list[list[str]]):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    tbl = Table(data, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16212c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f4f4")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tbl


def render_pdf(
    output_path: str,
    document_info: dict[str, Any],
    analytics: dict[str, Any],
    narrative: dict[str, str],
    narrative_source: str,
    charts: dict[str, str],
    sources: list[dict[str, Any]],
):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=8)
    body = styles["BodyText"]
    small = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=8, textColor=rl_colors.grey)

    doc = SimpleDocTemplate(output_path, pagesize=LETTER,
                             leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                             topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    story = []

    story.append(Paragraph("Indian Coal Mines — Data Intelligence Report", h1))
    story.append(Paragraph(
        f"Dataset period: {analytics['dataset_period']} | "
        f"Source document: {document_info.get('filename', 'N/A')} | "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        small,
    ))
    story.append(Spacer(1, 14))

    story.append(Paragraph("1. Executive Summary", h2))
    story.append(Paragraph(narrative["executive_summary"], body))

    story.append(Paragraph("2. Dataset Overview", h2))
    t = analytics["totals"]
    overview_data = [
        ["Metric", "Value"],
        ["Source file", document_info.get("filename", "N/A")],
        ["Sheet", document_info.get("sheet_name", "N/A")],
        ["Total mine records", str(t["mine_count"])],
        ["Mines with reported production", str(t["mines_with_production"])],
        ["Mines missing production value", str(t["mines_missing_production"])],
        ["Dataset period", analytics["dataset_period"]],
    ]
    story.append(_pdf_table(overview_data))

    story.append(Paragraph("3. Production Analysis", h2))
    story.append(Paragraph(narrative["production_analysis"], body))
    prod_data = [["Metric", "Value (MT)"],
                 ["Total production", f"{t['total_production_mt']:,}"],
                 ["Average per mine", f"{t['avg_production_mt']:,}"],
                 ["Maximum (single mine)", f"{t['max_production_mt']:,}"],
                 ["Minimum (single mine)", f"{t['min_production_mt']:,}"]]
    story.append(_pdf_table(prod_data))
    if "by_state" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["by_state"], width=6 * inch, height=3.4 * inch))
    if "coal_vs_lignite" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["coal_vs_lignite"], width=3.5 * inch, height=3.5 * inch))
    if "mine_type" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["mine_type"], width=5 * inch, height=3.3 * inch))
    if "ownership" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["ownership"], width=5 * inch, height=3.3 * inch))

    story.append(PageBreak())
    story.append(Paragraph("4. Mine / Owner Analysis", h2))
    story.append(Paragraph(narrative["mine_owner_analysis"], body))
    if "top_mines" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["top_mines"], width=6 * inch, height=3.4 * inch))
    top_mines_data = [["Mine", "State", "Owner", "Production (MT)"]]
    for m in analytics["top_mines"][:10]:
        top_mines_data.append([m["mine_name"][:35], m["state_ut"] or "-", m["owner_short"] or "-",
                                f"{m['production_mt']:,}"])
    story.append(_pdf_table(top_mines_data))

    if "by_owner" in charts:
        story.append(Spacer(1, 10))
        story.append(Image(charts["by_owner"], width=6 * inch, height=3.4 * inch))

    story.append(PageBreak())
    story.append(Paragraph("5. Key Findings", h2))
    story.append(Paragraph(narrative["key_findings"], body))

    story.append(Paragraph("6. Data Quality", h2))
    story.append(Paragraph(narrative["data_quality"], body))
    if "validation" in charts:
        story.append(Spacer(1, 8))
        story.append(Image(charts["validation"], width=5 * inch, height=3.3 * inch))
    v = analytics["validation"]
    val_data = [["Status", "Count"]] + [[k, str(vv)] for k, vv in v["by_status"].items()]
    story.append(_pdf_table(val_data))

    story.append(Paragraph("7. Important Observations", h2))
    story.append(Paragraph(narrative["important_observations"], body))

    story.append(Paragraph("8. Tables", h2))
    story.append(Paragraph("Production by State (top 10)", styles["Heading3"]))
    state_data = [["State/UT", "Mine Count", "Production (MT)"]]
    for r in analytics["by_state"][:10]:
        state_data.append([r["state_ut"], str(r["mine_count"]), f"{r['total_production_mt']:,}"])
    story.append(_pdf_table(state_data))

    story.append(Paragraph("9. Charts", h2))
    story.append(Paragraph(
        "All charts above are generated directly from the PostgreSQL analytical results; "
        "no chart is rendered from placeholder or assumed data.",
        body,
    ))

    story.append(Paragraph("10. Sources", h2))
    if sources:
        for s in sources[:20]:
            label = s.get("source_label") or "Source"
            url = s.get("source_url")
            if url:
                story.append(Paragraph(f'{label}: <link href="{url}">{url}</link>', small))
            else:
                story.append(Paragraph(label, small))
    else:
        story.append(Paragraph("No source URLs were recorded in the source document for this dataset.", body))

    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Narrative generated by: {narrative_source}", small))

    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# DOCX assembly (python-docx)
# ---------------------------------------------------------------------------
def _docx_table(doc, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def render_docx(
    output_path: str,
    document_info: dict[str, Any],
    analytics: dict[str, Any],
    narrative: dict[str, str],
    narrative_source: str,
    charts: dict[str, str],
    sources: list[dict[str, Any]],
):
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    doc.add_heading("Indian Coal Mines — Data Intelligence Report", level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Dataset period: {analytics['dataset_period']}  |  "
        f"Source document: {document_info.get('filename', 'N/A')}  |  "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).italic = True

    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(narrative["executive_summary"])

    doc.add_heading("2. Dataset Overview", level=2)
    t = analytics["totals"]
    _docx_table(doc, ["Metric", "Value"], [
        ["Source file", document_info.get("filename", "N/A")],
        ["Sheet", document_info.get("sheet_name", "N/A")],
        ["Total mine records", str(t["mine_count"])],
        ["Mines with reported production", str(t["mines_with_production"])],
        ["Mines missing production value", str(t["mines_missing_production"])],
        ["Dataset period", analytics["dataset_period"]],
    ])

    doc.add_heading("3. Production Analysis", level=2)
    doc.add_paragraph(narrative["production_analysis"])
    _docx_table(doc, ["Metric", "Value (MT)"], [
        ["Total production", f"{t['total_production_mt']:,}"],
        ["Average per mine", f"{t['avg_production_mt']:,}"],
        ["Maximum (single mine)", f"{t['max_production_mt']:,}"],
        ["Minimum (single mine)", f"{t['min_production_mt']:,}"],
    ])
    for key in ("by_state", "coal_vs_lignite", "mine_type", "ownership"):
        if key in charts:
            doc.add_picture(charts[key], width=Inches(5.5))

    doc.add_page_break()
    doc.add_heading("4. Mine / Owner Analysis", level=2)
    doc.add_paragraph(narrative["mine_owner_analysis"])
    if "top_mines" in charts:
        doc.add_picture(charts["top_mines"], width=Inches(5.5))
    _docx_table(
        doc, ["Mine", "State", "Owner", "Production (MT)"],
        [[m["mine_name"][:35], m["state_ut"] or "-", m["owner_short"] or "-", f"{m['production_mt']:,}"]
         for m in analytics["top_mines"][:10]],
    )
    if "by_owner" in charts:
        doc.add_picture(charts["by_owner"], width=Inches(5.5))

    doc.add_page_break()
    doc.add_heading("5. Key Findings", level=2)
    doc.add_paragraph(narrative["key_findings"])

    doc.add_heading("6. Data Quality", level=2)
    doc.add_paragraph(narrative["data_quality"])
    if "validation" in charts:
        doc.add_picture(charts["validation"], width=Inches(5))
    v = analytics["validation"]
    _docx_table(doc, ["Status", "Count"], [[k, str(vv)] for k, vv in v["by_status"].items()])

    doc.add_heading("7. Important Observations", level=2)
    doc.add_paragraph(narrative["important_observations"])

    doc.add_heading("8. Tables", level=2)
    doc.add_paragraph("Production by State (top 10)")
    _docx_table(
        doc, ["State/UT", "Mine Count", "Production (MT)"],
        [[r["state_ut"], str(r["mine_count"]), f"{r['total_production_mt']:,}"] for r in analytics["by_state"][:10]],
    )

    doc.add_heading("9. Charts", level=2)
    doc.add_paragraph(
        "All charts above are generated directly from the PostgreSQL analytical results; "
        "no chart is rendered from placeholder or assumed data."
    )

    doc.add_heading("10. Sources", level=2)
    if sources:
        for s in sources[:20]:
            label = s.get("source_label") or "Source"
            url = s.get("source_url")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}: {url}" if url else label)
    else:
        doc.add_paragraph("No source URLs were recorded in the source document for this dataset.")

    footer = doc.add_paragraph()
    footer.add_run(f"Narrative generated by: {narrative_source}").italic = True

    doc.save(output_path)
    return output_path
